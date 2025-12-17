import streamlit as st
import anthropic
import base64
import pandas as pd
import os
import zipfile
import time
import re
from docx import Document
from io import BytesIO

# --- 1. PAGE SETUP (MUST BE FIRST) ---
st.set_page_config(
    page_title="Pre-IB Lab Grader", 
    page_icon="🧪", 
    layout="wide"
)

# --- 2. CONFIGURATION & SECRETS ---
if "ANTHROPIC_API_KEY" in st.secrets:
    API_KEY = st.secrets["ANTHROPIC_API_KEY"]
elif "ANTHROPIC_API_KEY" in os.environ:
    API_KEY = os.environ.get("ANTHROPIC_API_KEY")
else:
    st.error("🚨 API Key not found!")
    st.info("On Streamlit Cloud, add your key to the 'Secrets' settings.")
    st.stop()

# --- 3. HARDCODED RUBRIC ---
PRE_IB_RUBRIC = """TOTAL: 100 POINTS (10 pts per section)

1. FORMATTING (10 pts):
- Criteria: Third-person passive voice, professional tone, superscripts/subscripts used correctly.
- DEDUCTIONS: 1-2 subscript errors = -0.5 pts. 3+ errors = -1.0 pt.

2. INTRODUCTION (10 pts):
- Criteria: Clear objective, background theory, balanced equations.
- OBJECTIVE: Must be explicit. (Missing: -1.0. Present but Vague/Implicit: -0.5).
- EQUATION: Balanced chemical equation required. (Missing: -1.0).
- THEORY/BACKGROUND: Must be thorough and connected to objective. (Irrelevant/Missing: -1.0. Brief/Not thoroughly connected: -0.5).
- NOTE: Do NOT deduct for inconsistent temperature units or citation context.

3. HYPOTHESIS (10 pts):
- Criteria: Specific prediction with scientific justification.
- JUSTIFICATION: Scientific reasoning required. (Missing: -2.0. Incomplete/Vague: -1.0).
- UNITS: Must include units for BOTH IV and DV. (Missing: -1.0, Incomplete: -0.5).
- MEASUREMENT: Specific description of how DV is measured. (Missing: -1.0, Vague: -0.5).

4. VARIABLES (10 pts):
- Criteria: IV, DV, 3+ Controls.
- SCORING: 
  * 10/10: All defined + explanations.
  * 9.5/10: DV measurement vague (-0.5).
  * 9.0/10: Explanations missing (-1.0).
  * 7.0/10: IV and DV variables missing (-3.0)
  * 8.0/10: IV or DV variable missing (-2.0)
  * 6.0/10: Control variables missing (-4.0)
  * 8/10: Only 2 control variables given and described (-2.0)
  * 9/10: All control varialbes not justified (-1.0)
  * 9.5/10: Justification of control variables vague (-0.5)

5. PROCEDURES (10 pts):
- Criteria: Numbered steps, quantities, safety.
- DIAGRAM: Diagram or photograph of experimental setup required. (Missing: -0.5).

6. RAW DATA (10 pts):
- Criteria: Qualitative observations, tables, units, sig figs.

7. DATA ANALYSIS (10 pts):
- Criteria: Calculation shown, Graph (Scatterplot, Trendline, Equation, R^2).
- GRAPH EQUATION: Linear equation must be displayed on graph. (Missing: -1.0).
- GRAPH R²: R² value must be displayed on graph. (Missing: -1.0).
- CALCULATIONS: Must be detailed and clear. (Unclear: -1.0).
- CALCULATION STEPS: All steps must be clearly explained OR labeled for clarity. (Not done: -0.5).
- NOTE: Intermediate precision allowed. Check final answer sig figs.

8. CONCLUSION (10 pts) [STRICT DEDUCTIONS]:
- HYPOTHESIS SUPPORT: Must indicate if data supports hypothesis. (If missing: -1.0).
- OUTLIERS/OMISSIONS: Must address data outliers or omissions. (No mention: -1.0. Mentioned but vague: -0.5).
- IV/DV RELATIONSHIP: Must explain graph trend. (If poor: -1.0).
- THEORY: Connect to chemical theory. (If missing: -1.0).
- QUANTITATIVE SUPPORT: Must cite specific numbers. (If missing: -2.0).
- QUALITATIVE SUPPORT: Must cite observations. (If missing: -0.5).
- LITERATURE COMPARISON: If comparison to literature is vague (no specific values), -0.5 pt.
* **Statistics (R vs R² CHECK):**
        * **R (Correlation):** * **Is the R value listed?** -> If NO, deduct 1.0.
            * **Is the explanation valid?** -> If the explanation is vague OR the student confuses R with R² (e.g., "The R² shows a positive correlation"), deduct 0.5.
        * **R² (Determination):** Must explain % variation/fit. (Missing entirely -> -1.0. Vague explanation -> -0.5).
- NOTE: Do NOT deduct for "Internal Inconsistency" or Citations here.

9. EVALUATION (10 pts) [STRICT QUALITY GATES]:
- REQUIREMENT: List errors + Specific Directional Impact on Data + Specific Improvement.
- ERROR CLASSIFICATION: Check if student uses terms "Systematic" or "Random". (If both terms are missing: -0.5. If present, NO deduction).
- QUANTITATIVE IMPACT SCORING (CRITICAL):
  * Requirement: For EVERY listed error, the student must state exactly how it changed the final calculated value (e.g., "This caused the calculated molar mass to be too high").
  * 0 Impact Descriptions: Deduct 2.0 pts (Score 8.0 max).
  * Some (but not all) Impact Descriptions: Deduct 1.0 pt (Score 9.0 max).
  * All Impact Descriptions Present: No deduction.
- IMPROVEMENT SCORING:
  * Specific equipment named = No deduction.
  * Vague ("use better scale") = Deduct 0.5.
  * Generic ("be careful") = Deduct 2.0.

10. REFERENCES (10 pts):
- 1 Reference only: -5.0 pts (Score 5.0).
- 2 References only: -3.0 pts (Score 7.0).
- 3+ References: Base score 10.0.
- FORMATTING: If APA formatting attempted but contains major errors, deduct 0.5 pts.
"""

# --- 4. SYSTEM PROMPT ---
SYSTEM_PROMPT = """You are an expert Pre-IB Chemistry Lab Grader. 
Your goal is to grade student lab reports according to the specific rules below.

### 🧠 FEEDBACK QUALITY STANDARDS (CRITICAL):
1.  **STRENGTHS (COMPREHENSIVE):** * Do not give generic praise (e.g., "Good job"). 
    * **Requirement:** You must summarize exactly *what* the student did well, **QUOTE** the specific text from their report that demonstrates this strength, and explain *why* it meets the rubric standard.
2.  **IMPROVEMENTS (ACTIONABLE):** * Do not just list the error. 
    * **Requirement:** For every deduction, you must provide:
        * **The Error:** What they wrote (or what was missing).
        * **The Fix:** A specific example of how to rewrite it or what to add.
        * **The Reason:** Why this is required by the rubric.

### ⚖️ CONSISTENCY PROTOCOL (MANDATORY):
1. **NO CURVING:** Grade every student exactly against the rubric. Do not compare students to each other.
2. **ISOLATED EVALUATION:** If a requirement is missing, deduct the points immediately. Do not "give credit" because the rest of the report was good.
3. **RIGID ADHERENCE:** Use the exact deduction values listed below. Do not approximate.

### ⚖️ CALIBRATION & TIE-BREAKER STANDARDS (MUST FOLLOW):

1.  **THE "BENEFIT OF DOUBT" RULE:**
    * If a student's phrasing is clumsy but technically accurate -> **NO DEDUCTION.**
    * If a student uses the wrong vocabulary word but the concept is correct -> **-0.5 (Vague).**
    * If the text is contradictory (says X, then says Not X) -> **-1.0 (Unclear).**

2.  **THE "DOUBLE JEOPARDY" BAN:**
    * Do NOT deduct points for the same error in two different sections.
    * *Example:* If they miss the units in the *Raw Data* table, deduct there. Do NOT also deduct for "missing units" in the *Analysis* section unless they made a *new* error there.

3.  **THE "STRICT BINARY" DECISION TREE:**
    * **Is the Hypothesis Justification missing?** * YES -> -2.0.
        * NO, but it relies on non-scientific reasoning (e.g., "I feel like...") -> -1.0.
    * **Is the R² value on the graph?**
        * YES (Explicitly written) -> 0 deduction.
        * NO (Not visible) -> -1.0 deduction. (Do not assume it is "implied").

4.  **IMAGE/TEXT CONFLICT:**
    * If the text says one thing (e.g., "R² = 0.98") but the graph image shows another (e.g., "R² = 0.50") -> **Trust the Image** and deduct for the discrepancy.
### 🧠 SCORING ALGORITHMS (STRICT ENFORCEMENT):

**CRITICAL INSTRUCTION:** 1. Perform ALL math calculations for ALL sections inside a single `<math_scratchpad>` block at the VERY START of your response. 
2. The user will NOT see this block (it is filtered out).
3. Do NOT include any math or deduction logic in the "OUTPUT FORMAT" sections. Only the final feedback text.

1.  **INTRODUCTION (Section 2) - DEDUCTION PROTOCOL:**
    * **Start at 10.0 Points.**
    * **Objective:** If Missing -> -1.0. If Vague/Implicit -> -0.5.
    * **Chemical Equation:** If Missing -> -1.0.
    * **Background Theory:** If Missing/Irrelevant -> -1.0. If Brief or Not thoroughly connected to objective -> -0.5.
    * **RESTRICTIONS (Do NOT Deduct):** No deductions for citation context or inconsistent units.

2.  **CONCLUSION (Section 8) - STRICT MATH PROTOCOL:**
    * **Start at 10.0 Points.**
    * **Hypothesis Support:** Not stated? -> -1.0.
    * **Outliers/Omissions:** No mention? -> -1.0. Vague? -> -0.5.
    * **Literature Comparison:** Vague comparison (no specific values)? -> -0.5.
    * **IV/DV Trend:** Missing logic? -> -1.0.
    * **Quantitative Data:** No numbers quoted? -> -2.0.
    * **Theory:** No connection? -> -1.0.
    * **Statistics (R vs R² CHECK):**
        * **R (Correlation):** Must explain Strength AND Direction. (Missing/No explanation -> -1.0. Vague explanation -> -0.5).
        * **R² (Determination):** Must explain % variation/fit. (Missing entirely -> -2.0. Vague explanation -> -1.0).
        * **Differentiation:** Ensure student treats R and R² as separate concepts. If they mix them up, apply the "Vague" deduction for both.
    * **Focus:** Repetitive/Unfocused? -> -0.5 (Max).
    * **RESTRICTIONS (Do NOT Deduct):** NO deductions for Citations, "Internal Inconsistency", or "Data Reliability".

3.  **HYPOTHESIS (Section 3):**
    * **Justification Check:** Missing? -> -2.0. Incomplete/Vague? -> -1.0.
    * **Units Check:** Missing -> -1.0. Incomplete -> -0.5.
    * **Measurement Check:** Missing -> -1.0. Vague -> -0.5.

4.  **VARIABLES (Section 4) - JUSTIFICATION PROTOCOL:**
    * **Control Justification:** * No justification given for why controls were chosen? -> -1.0.
        * Partial/Vague justification? -> -0.5.
    * **DV Measurement:** Method for measuring DV is vague? -> -0.5.
    * **Identification:** Any variable (IV, DV, Controls) missing? -> -1.0 per missing item.

5.  **DATA ANALYSIS (Section 7):**
    * **Trendline Equation:** Not shown on graph? -> -1.0.
    * **R² Value:** Not shown on graph? -> -1.0.
    * **Calculations:** Example calculations unclear? -> -1.0.
    * **Steps:** Calculation steps not clearly explained OR labeled? -> -0.5.

6.  **PROCEDURES (Section 5):**
    * **Diagram Check:** Diagram or photograph of experimental setup missing? -> -0.5.

7.  **EVALUATION (STRICT IMPACT AUDIT):** - **ERROR CLASSIFICATION (KEYWORD SEARCH):** Scan the text for the words "Systematic" or "Random". 
     * **If present:** Assume the student has differentiated correctly. DO NOT DEDUCT.
     * **If absent:** Deduct 0.5.
   - **MANDATORY IMPACT CHECK:** List every error the student mentions. For EACH error, verify if they explain 
     the DIRECTIONAL impact on the final calculated value (e.g., 'caused molar mass to be too high', 
     'made concentration lower than actual'). 
   - **SCORING:** If 0 errors have directional impact -> -2.0 pts. If some but not all -> -1.0 pt. 
     If all errors have direction -> No deduction.
   - In your feedback, you MUST write: 'You listed [X] errors. [Y] had explicit directional impact.' 
   - Penalize vague improvements (-0.5) or generic improvements like 'be more careful' (-2.0)."
    * **CRITICAL IMPACT AUDIT (THE "DIRECTION" CHECK - STRICTLY ENFORCE):**
        * **Step 1:** Count the TOTAL number of errors the student lists (e.g., "spilling water", "heat loss", "scale precision").
        * **Step 2:** For EACH error, search for EXPLICIT directional language about the calculated result:
            - ACCEPTABLE phrases: "made the result too high", "caused an overestimation", "led to a lower value", "increased the calculated mass", "decreased the final answer"
            - NOT ACCEPTABLE: "affected accuracy", "caused error", "impacted results", "reduced precision" (these are vague - no direction specified)
        * **Step 3:** Count how many errors have explicit directional impact.
        * **Step 4:** Apply Scoring (NO EXCEPTIONS):
            - If **ZERO** errors have directional impact explained -> **DEDUCT 2.0 points** (Max score 8.0)
            - If **SOME BUT NOT ALL** errors have directional impact -> **DEDUCT 1.0 point** (Max score 9.0)
            - If **ALL** errors have specific directional impact -> **NO DEDUCTION** (Score 10.0 possible)
        
        * **EXAMPLE GRADING:**
            - Student lists 3 errors but only explains direction for 2 of them -> DEDUCT 1.0 pt
            - Student lists 4 errors but explains direction for 0 of them -> DEDUCT 2.0 pts
            - Student lists 2 errors and explains direction for both -> NO DEDUCTION (assuming other criteria met)
    
    * **IMPROVEMENTS:** Specific equipment named? -> No deduction. Vague? -> -0.5. Generic? -> -2.0.
    
    * **MANDATORY FEEDBACK FORMAT:** In your response, you MUST explicitly state:
        - "You listed [X] total errors."
        - "Of these, [Y] had explicit directional impact on the calculated value."
        - If Y < X: "This results in a deduction of [1.0 or 2.0] points."

8.  **REFERENCES (Section 10) - QUANTITY CHECK:**
    * 1 Reference: Max Score 5.0.
    * 2 References: Max Score 7.0.
    * 3+ References: Max Score 10.0.
    * **Formatting:** Do NOT deduct for minor APA formatting errors. Only deduct for major errors.

### 📝 FEEDBACK STYLE INSTRUCTIONS:
1. **FORMATTING:** Use <sub> and <sup> tags for chemical formulas and exponents (e.g., write H<sub>2</sub>O, 10<sup>5</sup>).
2. **AVOID ROBOTIC CHECKLISTS:** Do not use "[Yes/No]".
3. **EXPLAIN WHY:** Write 2-3 sentences for each section.
4. **TOP 3 ACTIONABLE STEPS:** You MUST provide exactly THREE specific, actionable steps at the end. These should be concrete recommendations the student can implement in their next lab report.

### OUTPUT FORMAT:
Please strictly use the following format. Do not use horizontal rules (---) between sections. Do NOT print the calculation steps here.

# 📝 SCORE: [Total Points]/100
STUDENT: [Filename]

**📊 OVERALL SUMMARY & VISUAL ANALYSIS:**
* [1-2 sentences on quality]
* [Critique of graphs/images]

**📝 DETAILED RUBRIC BREAKDOWN:**

**1. FORMATTING: [Score]/10**
* **✅ Strengths:** [Detailed explanation of tone/voice quality]
* **⚠️ Improvements:** [**MANDATORY:** "Found [X] subscript errors." (If X=1 or 2, Score **MUST** be 9.5. If X>=3, Score is 9.0 or lower).]

**2. INTRODUCTION: [Score]/10**
* **✅ Strengths:** [Detailed explanation of objective/theory coverage]
* **⚠️ Improvements:** [**CRITICAL CHECKS:** * "Objective explicit?" (-1.0 if No, -0.5 if Vague). * "Chemical Equation present?" (-1.0 if No). * "Background thoroughly explained?" (-1.0 if No, -0.5 if Brief or not connected to objective). NOTE: Do not penalize citation context or unit consistency.]

**3. HYPOTHESIS: [Score]/10**
* **✅ Strengths:** [Quote prediction and praise the scientific reasoning]
* **⚠️ Improvements:** [**CRITICAL CHECKS:**
* "Justification: [Present/Missing/Vague]" (-2.0 if missing, -1.0 if vague/incomplete).
* "Units for IV/DV: [Present/Missing]" (-1.0 if missing, -0.5 if partial).
* "DV Measurement Description: [Specific/Vague/Missing]" (-1.0 if missing, -0.5 if vague).]

**4. VARIABLES: [Score]/10**
* **✅ Strengths:** [**LIST:** "Identified IV: [X], DV: [Y], Controls: [A, B, C]" and comment on clarity.]
* **⚠️ Improvements:** [If DV measurement is vague, state: "The method for measuring the DV was vague (-0.5 pts)." Suggest specific improvement.]

**5. PROCEDURES: [Score]/10**
* **✅ Strengths:** [Comment on reproducibility and safety details]
* **⚠️ Improvements:** [**DIAGRAM CHECK:** "Diagram of experimental setup included?" (-0.5 if missing). Identify exactly which step is vague and how to fix it.]

**6. RAW DATA: [Score]/10**
* **✅ Strengths:** [Comment on data organization and unit clarity]
* **⚠️ Improvements:** [Quote values with wrong units/sig figs and explain the correct format. Comment on inconsistent sig fig reporting for measuring tools.]

**7. DATA ANALYSIS: [Score]/10**
* **✅ Strengths:** [Summarize the calculation process. If Graph is perfect, mention that the scatterplot, equation, and labels are all correct here.]
* **⚠️ Improvements:** [**GRAPH AUDIT:** "Trendline Equation: [Present/Missing]" (-1.0 if missing). "R² Value: [Present/Missing]" (-1.0 if missing).
**CALCULATION AUDIT:** "Example calculations were [Clear/Unclear]." (If unclear, -1.0 pts). "Calculation steps were [Clearly Explained/Not Labeled or Explained]." (If not labeled/explained, -0.5 pts).]

**8. CONCLUSION (10 pts) [STRICT DEDUCTIONS]:
- HYPOTHESIS SUPPORT: Must indicate if data supports hypothesis. (If missing: -1.0).
- OUTLIERS/OMISSIONS: Must address data outliers or omissions. (No mention: -1.0. Mentioned but vague: -0.5).
- IV/DV RELATIONSHIP: Must explain graph trend. (If poor: -1.0).
- THEORY: Connect to chemical theory. (If missing: -1.0).
- QUANTITATIVE SUPPORT: Must cite specific numbers. (If missing: -2.0).
- QUALITATIVE SUPPORT: Must cite observations. (If missing: -0.5).
- LITERATURE COMPARISON: If comparison to literature is vague (no specific values), -0.5 pt.
- STATISTICS (CORRELATION COEFFICIENT - R):
  * Requirement: Must explicitly list the R value. (Missing: -1.0).
  * Explanation: Must explain Strength & Direction.
  * DEDUCTION: If R is present but explanation is vague OR student confuses R with R² (e.g., uses R² to describe direction) = -0.5 pts.
- STATISTICS (R² - DETERMINATION):
  * Requirement: Must explain Fit/Variability. (Missing: -1.0. Vague: -0.5).
- NOTE: Do NOT deduct for "Internal Inconsistency" or Citations here.

**9. EVALUATION: [Score]/10**
* **✅ Strengths:** [**LIST:** "You identified: [Error 1], [Error 2]..." and comment on depth.]
* **⚠️ Improvements:** [**ERROR CLASSIFICATION:** "You did not differentiate between systematic and random errors. (-0.5 pt)" OR "You successfully distinguished systematic from random errors."
**IMPACT/IMPROVEMENT AUDIT:** * "You listed [X] errors but only provided specific directional impacts for [Y] of them. (-1 pt)"
  * "Improvements were listed but were slightly vague (e.g., did not name specific equipment). (-0.5 pt)" ]

**10. REFERENCES: [Score]/10**
* **✅ Strengths:** [**MANDATORY:** "Counted [X] credible sources."]
* **⚠️ Improvements:** [**QUANTITY CHECK:** "Only found [X] sources." (If 1 source -> Score 5.0. If 2 sources -> Score 7.0). **FORMATTING:** "APA Formatting Check: [Correct/Incorrect]" (-0.5 if incorrect).]

**💡 TOP 3 ACTIONABLE STEPS FOR NEXT TIME:**
1. [Step 1 - Specific and concrete recommendation]
2. [Step 2 - Specific and concrete recommendation]
3. [Step 3 - Specific and concrete recommendation]
"""

# --- 5. SESSION STATE INITIALIZATION ---
if 'autosave_dir' not in st.session_state:
    # 1. Initialize Autosave Folder (The Fix from before)
    base_folder = "autosave_feedback_pre-ib"
    current_dir = os.getcwd()
    full_path = os.path.join(current_dir, base_folder)
    
    if not os.path.exists(full_path):
        os.makedirs(full_path)
        print(f"📁 Created autosave folder at: {full_path}")
    
    st.session_state.autosave_dir = full_path

# 2. Initialize other required variables (Restoring these fixes the error)
if 'current_results' not in st.session_state:
    st.session_state.current_results = []

if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False

if 'current_session_name' not in st.session_state:
    st.session_state.current_session_name = f"Session_{time.strftime('%H%M')}"

if 'saved_sessions' not in st.session_state:
    st.session_state.saved_sessions = {}

# Debug Helper
st.sidebar.success(f"📂 Autosave Folder: `{st.session_state.autosave_dir}`")

client = anthropic.Anthropic(api_key=API_KEY)

# --- 6. HELPER FUNCTIONS ---
def encode_file(uploaded_file):
    try:
        uploaded_file.seek(0)
        return base64.b64encode(uploaded_file.read()).decode('utf-8')
    except Exception as e:
        st.error(f"Error encoding file: {e}")
        return None

def get_media_type(filename):
    ext = filename.lower().split('.')[-1]
    media_types = {
        'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
        'gif': 'image/gif', 'webp': 'image/webp', 'pdf': 'application/pdf'
    }
    return media_types.get(ext, 'image/jpeg')

# --- UPDATED TEXT EXTRACTION (WITH SUBSCRIPT DETECTION) ---
def get_para_text_with_formatting(para):
    """Iterate through runs to capture subscript/superscript formatting."""
    text_parts = []
    for run in para.runs:
        text = run.text
        # Check for subscript
        if run.font.subscript:
            text = f"<sub>{text}</sub>"
        # Check for superscript
        elif run.font.superscript:
            text = f"<sup>{text}</sup>"
        text_parts.append(text)
    return "".join(text_parts)

def extract_text_from_docx(file):
    try:
        file.seek(0) 
        doc = Document(file)
        full_text = []
        
        # 1. Extract Paragraphs with Formatting
        for para in doc.paragraphs:
            full_text.append(get_para_text_with_formatting(para))
            
        # 2. Extract Tables with Formatting
        if doc.tables:
            full_text.append("\n--- DETECTED TABLES ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Extract paragraphs within cell
                        cell_content = []
                        for para in cell.paragraphs:
                            cell_content.append(get_para_text_with_formatting(para))
                        row_text.append(" ".join(cell_content).strip())
                    full_text.append(" | ".join(row_text))
                full_text.append("\n") 
        
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading .docx file: {e}"

def extract_images_from_docx(file):
    images = []
    try:
        file.seek(0) # CRITICAL FIX: Reset pointer before reading
        with zipfile.ZipFile(file) as z:
            for filename in z.namelist():
                if filename.startswith('word/media/') and filename.split('.')[-1].lower() in ['png', 'jpg', 'jpeg', 'gif']:
                    img_data = z.read(filename)
                    b64_img = base64.b64encode(img_data).decode('utf-8')
                    ext = filename.split('.')[-1].lower()
                    images.append({
                        "type": "image",
                        "source": {
                            "type": "base64", 
                            "media_type": f"image/{'jpeg' if ext=='jpg' else ext}", 
                            "data": b64_img
                        }
                    })
    except Exception as e:
        print(f"Image extraction failed: {e}")
    return images

def process_uploaded_files(uploaded_files):
    final_files = []
    IGNORED_FILES = {'.ds_store', 'desktop.ini', 'thumbs.db', '__macosx'}
    VALID_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'gif', 'webp', 'docx'}
    
    file_counts = {"pdf": 0, "docx": 0, "image": 0, "ignored": 0}

    for file in uploaded_files:
        file_name_lower = file.name.lower()
        if file_name_lower in IGNORED_FILES or file_name_lower.startswith('._'):
            continue

        if file_name_lower.endswith('.zip'):
            try:
                with zipfile.ZipFile(file) as z:
                    for filename in z.namelist():
                        clean_name = filename.lower()
                        if any(x in clean_name for x in IGNORED_FILES) or filename.startswith('.'): continue
                        ext = clean_name.split('.')[-1]
                        if ext in VALID_EXTENSIONS:
                            file_bytes = z.read(filename)
                            virtual_file = BytesIO(file_bytes)
                            virtual_file.name = os.path.basename(filename)
                            final_files.append(virtual_file)
                            if ext == 'docx': file_counts['docx'] += 1
                            elif ext == 'pdf': file_counts['pdf'] += 1
                            else: file_counts['image'] += 1
            except Exception as e:
                st.error(f"Error unzipping {file.name}: {e}")
        else:
            ext = file_name_lower.split('.')[-1]
            if ext in VALID_EXTENSIONS:
                final_files.append(file)
                if ext == 'docx': file_counts['docx'] += 1
                elif ext == 'pdf': file_counts['pdf'] += 1
                else: file_counts['image'] += 1
            else:
                file_counts['ignored'] += 1
            
    return final_files, file_counts

# --- MATH CHECKER ---
def recalculate_total_score(text):
    try:
        pattern = r"\d+\.\s+[A-Z\s]+:\s+([\d\.]+)/10"
        matches = re.findall(pattern, text)
        if matches:
            total_score = sum(float(m) for m in matches)
            if total_score.is_integer():
                total_score = int(total_score)
            else:
                total_score = round(total_score, 1)
            # UPDATED REGEX FOR HEADER SCORE
            text = re.sub(r"#\s*📝\s*SCORE:\s*[\d\.]+/100", f"# 📝 SCORE: {total_score}/100", text, count=1)
    except Exception as e:
        print(f"Error recalculating score: {e}")
    return text

# --- IMPROVED CSV FEEDBACK PARSER ---
def parse_feedback_for_csv(text):
    data = {}
    
    # 1. Clean Textual Decorators
    clean_text = re.sub(r'[*#]', '', text) 
    
    # 2. Extract Overall Summary
    try:
        # Looks for "OVERALL SUMMARY" followed by text until "1. " or "DETAILED"
        summary_match = re.search(r"OVERALL SUMMARY.*?:\s*\n(.*?)(?=1\.|DETAILED)", clean_text, re.DOTALL | re.IGNORECASE)
        if summary_match:
            # AGGRESSIVE CLEANING: Collapse newlines to single space for CSV safety
            raw_summary = summary_match.group(1).strip()
            data["Overall Summary"] = re.sub(r'[\r\n]+', ' ', raw_summary)
        else:
            data["Overall Summary"] = "Summary not found"
    except Exception as e:
        data["Overall Summary"] = f"Parsing Error: {e}"

    # 3. Extract Section Scores and Comments
    # Regex looks for: "1. SECTION NAME: Score/10" followed by content
    sections = re.findall(r"(\d+)\.\s+([A-Za-z\s]+):\s+([\d\.]+)/10\s*\n(.*?)(?=\n\d+\.|\Z|💡)", clean_text, re.DOTALL)
    
    for _, name, score, content in sections:
        col_name = name.strip().title() # e.g. "Formatting"
        data[f"{col_name} Score"] = score
        
        # AGGRESSIVE CLEANING for CSV:
        # Replaces all whitespace (newlines, tabs) with a single space to prevent broken CSVs
        cleaned_feedback = re.sub(r'[\r\n]+', ' ', content.strip())
        data[f"{col_name} Feedback"] = cleaned_feedback

    return data

def clean_for_sheets(text):
    if not isinstance(text, str): return text
    text = re.sub(r'#+\s*', '', text)
    text = text.replace('**', '')
    return text.strip()

# --- NEW FUNCTION: CLEAN HIDDEN SCRATCHPAD ---
def clean_hidden_scratchpad(text):
    """Removes the internal <math_scratchpad> tags before displaying to the user."""
    return re.sub(r'<math_scratchpad>.*?</math_scratchpad>', '', text, flags=re.DOTALL | re.IGNORECASE).strip()

def grade_submission(file, model_id):
    ext = file.name.split('.')[-1].lower()
    
    if ext == 'docx':
        text_content = extract_text_from_docx(file)
        
        # Check for empty text to warn user
        if len(text_content.strip()) < 50:
            text_content += "\n\n[SYSTEM NOTE: Very little text text extracted. Content may be in images or text boxes.]"
            
        # Use String Concatenation instead of f-string to prevent brace errors
        prompt_text = (
            "Please grade this lab report based on the Pre-IB rubric below.\n"
            "Note: This is a converted Word Document. The text content is provided below, followed by any embedded images.\n\n"
            "⚠️ CRITICAL INSTRUCTIONS:\n"
            "1. **BE SPECIFIC & EXPANDED:** Write 2-3 sentences per section explaining the score. Quote text/data. No generic feedback.\n"
            "2. **VARIABLES:** List the exact variables found. If found, score 9-10.\n"
            "3. **REFERENCES:** Count the sources. If >= 3, MINIMUM score is 9.0.\n"
            "4. **FORMATTING MATH:** 1-2 errors = -0.5 pts (Score 9.5). 3+ errors = -1.0 pt (Score 9.0).\n"
            "5. **FORMATTING DETECTION:** The text has been pre-processed. Subscripts appear as <sub>text</sub>. Superscripts appear as <sup>text</sup>. If these tags are present, the student formatted it CORRECTLY. Do not penalize.\n"
            "6. **GRAPHS:** Check for R² (-1.0 if missing), Equation (-1.0 if missing), Scatterplot format, and Units. Place audit in Strengths if perfect.\n"
            "7. **CONCLUSION:** Check for Outliers/Omissions (-1.0 if not mentioned, -0.5 if vague), IV/DV trend (-1.0), Theory (-1.0), Quant Data (-2.0), Qual Data (-0.5). **R-VALUE CHECK:** Missing R value -> -1.0. Confuses R with R² OR Vague explanation -> -0.5. R² (-1.0 if missing, -0.5 if vague). Repetitiveness (-0.5).\n"
            "8. **DATA ANALYSIS:** Check calculations for clarity (-1.0 if unclear). Check if calculation steps are clearly explained or labeled (-0.5 if not). Do NOT penalize for missing uncertainty analysis.\n"
            "9. **EVALUATION:** Check if systematic vs random errors are differentiated (-0.5 if not). Penalize vague impact/improvements. Must specify DIRECTION of error and SPECIFIC equipment for **ALL** errors. (0 pts if missing, 1 pt if partial).\n"
            "10. **HYPOTHESIS:** Check Justification (-2.0 if missing, -1.0 if vague). Check Units for IV/DV (-1.0 if missing, -0.5 if incomplete). Check DV Measurement (-1.0 if missing, -0.5 if vague).\n"
            "11. **INTRODUCTION:** Check for Chemical Equation (-1.0 if missing). Check for Objective (-1.0 if missing, -0.5 if vague). Check Theory Relevance (-1.0 if irrelevant). Check if Theory connects to Objective (-0.5 if not thoroughly connected). Check Thoroughness (-1.0 if missing, -0.5 if brief). DO NOT penalize for inconsistent units. DO NOT penalize for citation context.\n"
            "12. **PROCEDURES:** Check if a diagram of the experimental setup is included (-0.5 if missing).\n"
            "13. **HIDDEN MATH:** Use <math_scratchpad> tags for all calculations.\n"
            "14. **COMPLETE RESPONSE:** Ensure all 10 sections are graded. Do not stop early.\n"
            "15. **TOP 3 ACTIONABLE STEPS:** You MUST provide exactly THREE specific, concrete, actionable recommendations at the end of your feedback.\n\n"
            "--- RUBRIC START ---\n" + PRE_IB_RUBRIC + "\n--- RUBRIC END ---\n\n"
            "STUDENT TEXT:\n" + text_content
        )
        
        user_message = [
            {
                "type": "text",
                "text": prompt_text
            }
        ]
        images = extract_images_from_docx(file)
        if images:
            user_message.extend(images)
    else:
        base64_data = encode_file(file)
        if not base64_data: return "Error processing file."
        media_type = get_media_type(file.name)
        
        prompt_text = (
            "Please grade this lab report based on the Pre-IB rubric below.\n\n"
            "--- RUBRIC START ---\n" + PRE_IB_RUBRIC + "\n--- RUBRIC END ---\n\n"
            "INSTRUCTIONS:\n"
            "1. **BE SPECIFIC & EXPANDED:** Write 2-3 sentences per section explaining the score. Quote text/data. No generic feedback.\n"
            "2. **VARIABLES:** List the exact variables found. If found, score 9-10.\n"
            "3. **REFERENCES:** Count the sources. If >= 3, MINIMUM score is 9.0.\n"
            "4. **FORMATTING MATH:** 1-2 errors = -0.5 pts (Score 9.5). 3+ errors = -1.0 pt (Score 9.0).\n"
            "5. **GRAPHS:** Check for R² (-1.0 if missing), Equation (-1.0 if missing), Scatterplot format, and Units. Place audit in Strengths if perfect.\n"
            "6. **CONCLUSION:** Check for Outliers/Omissions (-1.0 if not mentioned, -0.5 if vague), IV/DV trend (-1.0), Theory (-1.0), Quant Data (-2.0), Qual Data (-0.5), R Value (-1.0), R² (-1.0 if missing, -0.5 if vague), Repetitiveness (-0.5).\n"
            "7. **DATA ANALYSIS:** Check calculations for clarity (-1.0 if unclear). Check if calculation steps are clearly explained or labeled (-0.5 if not). Do NOT penalize for missing uncertainty analysis.\n"
            "8. **EVALUATION:** Check if systematic vs random errors are differentiated (-0.5 if not). Penalize vague impact/improvements. Must specify DIRECTION of error and SPECIFIC equipment for **ALL** errors. (0 pts if missing, 1 pt if partial).\n"
            "9. **HYPOTHESIS:** Check Justification (-2.0 if missing, -1.0 if vague). Check Units for IV/DV (-1.0 if missing, -0.5 if incomplete). Check DV Measurement (-1.0 if missing, -0.5 if vague).\n"
            "10. **INTRODUCTION:** Check for Chemical Equation (-1.0 if missing). Check for Objective (-1.0 if missing, -0.5 if vague). Check Theory Relevance (-1.0 if irrelevant). Check if Theory connects to Objective (-0.5 if not thoroughly connected). Check Thoroughness (-1.0 if missing, -0.5 if brief). DO NOT penalize for inconsistent units. DO NOT penalize for citation context.\n"
            "11. **PROCEDURES:** Check if a diagram of the experimental setup is included (-0.5 if missing).\n"
            "12. **HIDDEN MATH:** Use <math_scratchpad> tags for all calculations.\n"
            "13. **COMPLETE RESPONSE:** Ensure all 10 sections are graded. Do not stop early.\n"
            "14. **TOP 3 ACTIONABLE STEPS:** You MUST provide exactly THREE specific, concrete, actionable recommendations at the end of your feedback.\n"
        )
        
        user_message = [
            {
                "type": "text",
                "text": prompt_text
            },
            {
                "type": "document" if media_type == 'application/pdf' else "image",
                "source": {"type": "base64", "media_type": media_type, "data": base64_data}
            }
        ]

    max_retries = 5 
    retry_delay = 5 
    
    for attempt in range(max_retries):
        try:
            # Temperature=0 for Maximum Consistency
            response = client.messages.create(
                model=model_id, # Uses the ID passed from Sidebar
                max_tokens=4096, # MAX TOKEN LIMIT
                temperature=0.0,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_message}]
            )
            raw_text = response.content[0].text
            
            # --- CLEAN THE SCRATCHPAD ---
            cleaned_text = clean_hidden_scratchpad(raw_text)
            
            corrected_text = recalculate_total_score(cleaned_text)
            return corrected_text
            
        except (anthropic.RateLimitError, anthropic.APIStatusError) as e:
            # Check for Overloaded (529) or Rate Limit (429)
            if isinstance(e, anthropic.APIStatusError) and e.status_code == 529:
                status_msg = f"⚠️ Server Overloaded (529). Retrying attempt {attempt+1}/{max_retries}..."
                print(status_msg) # Log to console
                time.sleep(retry_delay * (attempt + 1)) # Exponential backoff
                continue
            
            if isinstance(e, anthropic.RateLimitError):
                status_msg = f"⚠️ Rate Limit Hit. Retrying attempt {attempt+1}/{max_retries}..."
                print(status_msg)
                time.sleep(retry_delay * (attempt + 1))
                continue
                
            return f"⚠️ Error: {str(e)}"
            
        except Exception as e:
            return f"⚠️ Error: {str(e)}"

# --- PARSE SCORE FUNCTION ---
def parse_score(text):
    """Extract the total score from Claude's feedback text."""
    try:
        match = re.search(r"#\s*📝\s*SCORE:\s*([\d\.]+)/100", text)
        if match:
            return match.group(1).strip()
        match = re.search(r"SCORE:\s*([\d\.]+)/100", text)
        if match:
            return match.group(1).strip()
    except Exception as e:
        print(f"Error parsing score: {e}")
    return "N/A"

# # --- WORD FORMATTER (Upgraded for Sub/Superscripts) ---
def write_markdown_to_docx(doc, text):
    """
    Parses Markdown text and writes it to a docx Document.
    Handles headers, bullet points, bold (**text**), 
    superscript (<sup>text</sup>), and subscript (<sub>text</sub>).
    """
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue 
        
        # 1. Handle Headers
        if line.startswith('# '): 
            doc.add_heading(line.replace('# ', '').replace('*', '').strip(), level=2) 
            continue
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', '').replace('*', '').strip(), level=3)
            continue
        if line.startswith('## '): 
            doc.add_heading(line.replace('## ', '').replace('*', '').strip(), level=2)
            continue
        if line.startswith('---') or line.startswith('___'):
            doc.add_paragraph("_" * 50) # visual separator
            continue

        # 2. Handle List Items
        if line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:] 
        else:
            p = doc.add_paragraph()
            content = line

        # 3. Handle Formatting Tags (Bold, Sup, Sub)
        # This regex splits the text by tags so we can process each chunk
        # It looks for **bold**, <sup>sup</sup>, and <sub>sub</sub>
        parts = re.split(r'(\*\*.*?\*\*|<sup>.*?</sup>|<sub>.*?</sub>)', content)
        
        for part in parts:
            if not part: continue # Skip empty splits
            
            run = p.add_run()
            
            # Handle Bold
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2].replace('<sub>', '').replace('</sub>', '').replace('<sup>', '').replace('</sup>', '')
                run.bold = True
            
            # Handle Superscript (Exponents)
            elif part.startswith('<sup>') and part.endswith('</sup>'):
                run.text = part[5:-6]
                run.font.superscript = True
                
            # Handle Subscript (Chemical Formulas)
            elif part.startswith('<sub>') and part.endswith('</sub>'):
                run.text = part[5:-6]
                run.font.subscript = True
                
            # Regular Text
            else:
                run.text = part


def create_master_doc(results, session_name):
    doc = Document()
    # REMOVED SESSION HEADER
    # doc.add_heading(f"Lab Report Grades: {session_name}", 0) 
    for item in results:
        # REMOVED FILENAME HEADER (Starts with Score + Student Name)
        write_markdown_to_docx(doc, item['Feedback'])
        doc.add_page_break()
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_zip_bundle(results):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as z:
        for item in results:
            doc = Document()
            # REMOVED FEEDBACK HEADER
            write_markdown_to_docx(doc, item['Feedback'])
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            safe_name = os.path.splitext(item['Filename'])[0] + "_Feedback.docx"
            z.writestr(safe_name, doc_buffer.getvalue())
    return zip_buffer.getvalue()

# --- NEW: AUTOSAVE INDIVIDUAL REPORT ---
def autosave_report(item, autosave_dir):
    """Save individual report as Word doc and append to CSV immediately after grading."""
    try:
        # --- FIX: FORCE FOLDER CREATION ---
        if not os.path.exists(autosave_dir):
            os.makedirs(autosave_dir)
        # ----------------------------------
        # 1. Save Word Document
        doc = Document()
        write_markdown_to_docx(doc, item['Feedback'])
        safe_filename = os.path.splitext(item['Filename'])[0] + "_Feedback.docx"
        doc_path = os.path.join(autosave_dir, safe_filename)
        doc.save(doc_path)
        
        # 2. Append to CSV (or create if doesn't exist)
        csv_path = os.path.join(autosave_dir, "gradebook.csv")
        
        # Parse feedback into row data
        row_data = {
            "Filename": item['Filename'],
            "Overall Score": item['Score']
        }
        feedback_data = parse_feedback_for_csv(item['Feedback'])
        row_data.update(feedback_data)
        
        # Check if CSV exists
        if os.path.exists(csv_path):
            existing_df = pd.read_csv(csv_path)
            # Remove duplicate if re-grading same file
            existing_df = existing_df[existing_df['Filename'] != item['Filename']]
            new_df = pd.concat([existing_df, pd.DataFrame([row_data])], ignore_index=True)
        else:
            new_df = pd.DataFrame([row_data])
        
        new_df.to_csv(csv_path, index=False, encoding='utf-8-sig')
        
        return True
    except Exception as e:
        print(f"Autosave failed for {item['Filename']}: {e}")
        return False
    
def display_results_ui():
    if not st.session_state.current_results:
        return

    st.divider()
    st.subheader(f"📊 Results: {st.session_state.current_session_name}")
    
    # --- PREPARE DATA ---
    results_list = []
    for item in st.session_state.current_results:
        row_data = {
            "Filename": item['Filename'],
            "Overall Score": item['Score']
        }
        feedback_data = parse_feedback_for_csv(item['Feedback'])
        row_data.update(feedback_data)
        results_list.append(row_data)
        
    csv_df = pd.DataFrame(results_list)
    
    # Sort columns
    cols = list(csv_df.columns)
    priority = ['Filename', 'Overall Score', 'Overall Summary']
    remaining = [c for c in cols if c not in priority]
    remaining.sort(key=lambda x: (x.split(' ')[0], 'Feedback' in x)) 
    final_cols = [c for c in priority if c in cols] + remaining
    csv_df = csv_df[final_cols]
    
    # --- DOWNLOADS ---
    csv_data = csv_df.to_csv(index=False).encode('utf-8-sig') 
    master_doc_data = create_master_doc(st.session_state.current_results, st.session_state.current_session_name)
    zip_data = create_zip_bundle(st.session_state.current_results)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("📄 Docs (.docx)", master_doc_data, f'{st.session_state.current_session_name}_Docs.docx', "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with col2:
        st.download_button("📦 Bundle (.zip)", zip_data, f'{st.session_state.current_session_name}_Students.zip', "application/zip", use_container_width=True)
    with col3:
        st.download_button("📊 CSV Export", csv_data, f'{st.session_state.current_session_name}_Detailed.csv', "text/csv", use_container_width=True)

    # --- AUTOSAVE INFO ---
    if os.path.exists(st.session_state.autosave_dir):
        st.caption(f"💾 Backup saved to: `{st.session_state.autosave_dir}`")

    # --- MAIN DISPLAY (RENDERED ONCE) ---
    st.divider()
    st.write("### 🏆 Gradebook")
    st.dataframe(csv_df, use_container_width=True)
    
    st.write("### 📝 Detailed Feedback History")
    # We use reversed() so the newest file is always at the top
    for idx, item in enumerate(reversed(st.session_state.current_results)):
        # Expand the very first item (newest), collapse others
        is_most_recent = (idx == 0)
        with st.expander(f"📄 {item['Filename']} (Score: {item['Score']}/100)", expanded=is_most_recent):
            st.markdown(item['Feedback'], unsafe_allow_html=True)

    # --- AUTOSAVE FOLDER ACCESS ---
    st.divider()
    st.info("💾 **Auto-saved files:** Individual feedback documents and gradebook are being saved to the `autosave_feedback_pre-ib` folder as grading progresses.")
    
    autosave_path = st.session_state.autosave_dir
    if os.path.exists(autosave_path):
        csv_autosave = os.path.join(autosave_path, "gradebook.csv")
        if os.path.exists(csv_autosave):
            with open(csv_autosave, 'rb') as f:
                st.download_button(
                    "📥 Download Auto-saved Gradebook (CSV)",
                    f.read(),
                    "autosaved_gradebook.csv",
                    "text/csv",
                    use_container_width=True
                )
        
# --- 6. SIDEBAR ---
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # UPDATED DEFAULT MODEL ID
    user_model_id = st.text_input(
        "🤖 Model ID", 
        value="claude-sonnet-4-20250514", 
        help="Change this if you have a specific Beta model or newer ID"
    )
    
    st.divider()
    st.header("💾 History Manager")
    save_name = st.text_input("Session Name", placeholder="e.g. Period 3 - Kinetics")
    if st.button("💾 Save Session"):
        if st.session_state.current_results:
            st.session_state.saved_sessions[save_name] = st.session_state.current_results
            st.success(f"Saved '{save_name}'!")
        else:
            st.warning("No results to save yet.")
            
    if st.session_state.saved_sessions:
        st.divider()
        st.subheader("📂 Load Session")
        selected_session = st.selectbox("Select Batch", list(st.session_state.saved_sessions.keys()))
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Load"):
                st.session_state.current_results = st.session_state.saved_sessions[selected_session]
                st.session_state.current_session_name = selected_session
                st.rerun()
        with col2:
            if st.button("🗑️ Delete"):
                del st.session_state.saved_sessions[selected_session]
                st.rerun()

    st.divider() 
    
    with st.expander("View Grading Criteria"):
        # CHANGED FROM st.text(PRE_IB_RUBRIC) TO st.text(IB_RUBRIC)
        st.text(PRE_IB_RUBRIC)

# --- 7. MAIN INTERFACE ---
st.title("🧪 Pre-IB Lab Grader")
st.caption(f"Current Session: **{st.session_state.current_session_name}**")

st.info("💡 **Tip:** To upload a folder, open it, press `Ctrl+A` (Select All), and drag everything here.")

raw_files = st.file_uploader(
    "📂 Upload Reports (PDF, Word, Images, ZIP)", 
    type=['pdf', 'docx', 'png', 'jpg', 'jpeg', 'zip'], 
    accept_multiple_files=True
)

processed_files = []
if raw_files:
    processed_files, counts = process_uploaded_files(raw_files)
    if len(processed_files) > 0:
        st.success(f"✅ Found **{len(processed_files)}** valid reports.")
        st.caption(f"📄 PDFs: {counts['pdf']} | 📝 Word Docs: {counts['docx']} | 🖼️ Images: {counts['image']}")
        if counts['ignored'] > 0:
            st.warning(f"⚠️ {counts['ignored']} files were ignored (unsupported format).")
    else:
        if raw_files:
            st.warning("No valid PDF, Word, or Image files found.")

# Find the grading button section (around line 820-890) and replace with this:

if st.button("🚀 Grade Reports", type="primary", disabled=not processed_files):
    
    st.write("---")
    progress = st.progress(0)
    status_text = st.empty()
    live_results_table = st.empty()
    
    # NEW: Placeholder for cumulative feedback display (cleared and rewritten each iteration)
    st.subheader("📋 Live Grading Feedback")
    feedback_placeholder = st.empty()
    
    # Initialize Session State list if not present
    if 'current_results' not in st.session_state:
        st.session_state.current_results = []
    
    # Create a set of already graded filenames for quick lookup
    existing_filenames = {item['Filename'] for item in st.session_state.current_results}
    
    for i, file in enumerate(processed_files):
        # 1. SMART RESUME CHECK: Skip if already graded
        if file.name in existing_filenames:
            status_text.info(f"↩ Skipping **{file.name}** (Already Graded)")
            time.sleep(0.5) # Brief pause for visual feedback
            progress.progress((i + 1) / len(processed_files))
            continue

        # 2. GRADING LOGIC
        status_text.markdown(f"**Grading:** `{file.name}` ({i+1}/{len(processed_files)})...")
        
        try:
            # Polite delay to prevent API overloading
            time.sleep(2) 
            
            feedback = grade_submission(file, user_model_id)
            score = parse_score(feedback)
            
            # 3. IMMEDIATE SAVE TO SESSION STATE
            new_entry = {
                "Filename": file.name,
                "Score": score,
                "Feedback": feedback
            }
            
            st.session_state.current_results.append(new_entry)
            
            # 4. AUTOSAVE TO DISK
            autosave_success = autosave_report(new_entry, st.session_state.autosave_dir)
            if autosave_success:
                status_text.success(f"✅ **{file.name}** graded & auto-saved! (Score: {score}/100)")
            else:
                status_text.warning(f"⚠️ **{file.name}** graded but autosave failed (Score: {score}/100)")
            
            # Update the existing set so duplicates within the same batch run are also caught
            existing_filenames.add(file.name)
            
            # 5. LIVE TABLE UPDATE
            df_live = pd.DataFrame(st.session_state.current_results)
            live_results_table.dataframe(df_live[["Filename", "Score"]], use_container_width=True)
            
            # 6. LIVE FEEDBACK DISPLAY (During grading only)
            with feedback_placeholder.container():
                for idx, item in enumerate(st.session_state.current_results):
                    is_most_recent = (idx == len(st.session_state.current_results) - 1)
                    with st.expander(f"📄 {item['Filename']} (Score: {item['Score']}/100)", expanded=is_most_recent):
                        st.markdown(item['Feedback'], unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"❌ Error grading {file.name}: {e}")
            
        progress.progress((i + 1) / len(processed_files))
        
    # 7. CLEAR LIVE GRADING DISPLAY AFTER COMPLETION
    status_text.success("✅ Grading Complete! All reports auto-saved.")
    progress.empty()
    feedback_placeholder.empty()  # ← THIS IS THE KEY FIX - Clears the live feedback
    live_results_table.empty()     # ← Also clear the live table
    
    # Show message about autosave location
    st.info(f"💾 **Backup Location:** All feedback has been saved to `{st.session_state.autosave_dir}/` folder. You can download individual files or the full gradebook below.")

# --- 8. PERSISTENT DISPLAY (This stays - it's called outside the grading loop) ---
if st.session_state.current_results:
    display_results_ui()